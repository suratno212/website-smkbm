from docx import Document
from docx.shared import Inches

doc = Document()
doc.add_heading('Struktur Tabel Database SIAKAD SMK', 0)

tabels = [
    {
        'title': 'users',
        'fields': [
            ('id', 'INT', '11', 'Primary Key, Auto Increment'),
            ('username', 'VARCHAR', '100', 'Username login'),
            ('password', 'VARCHAR', '255', 'Password hash'),
            ('role', 'ENUM', 'admin/guru/siswa', 'Role user'),
            ('created_at', 'DATETIME', '-', 'Tanggal dibuat'),
            ('updated_at', 'DATETIME', '-', 'Tanggal update'),
        ]
    },
    {
        'title': 'jurusan',
        'fields': [
            ('id', 'INT', '11', 'Primary Key, Auto Increment'),
            ('nama_jurusan', 'VARCHAR', '100', 'Nama jurusan'),
            ('created_at', 'DATETIME', '-', 'Tanggal dibuat'),
            ('updated_at', 'DATETIME', '-', 'Tanggal update'),
        ]
    },
    {
        'title': 'ruangan',
        'fields': [
            ('id', 'INT', '11', 'Primary Key, Auto Increment'),
            ('nama_ruangan', 'VARCHAR', '100', 'Nama ruangan'),
            ('created_at', 'DATETIME', '-', 'Tanggal dibuat'),
            ('updated_at', 'DATETIME', '-', 'Tanggal update'),
        ]
    },
    {
        'title': 'tahun_akademik',
        'fields': [
            ('id', 'INT', '11', 'Primary Key, Auto Increment'),
            ('tahun', 'VARCHAR', '9', 'Contoh: 2024/2025'),
            ('semester', 'ENUM', '1/2', 'Semester'),
            ('status', 'ENUM', 'Aktif/Tidak', 'Status tahun akademik'),
            ('created_at', 'DATETIME', '-', 'Tanggal dibuat'),
            ('updated_at', 'DATETIME', '-', 'Tanggal update'),
        ]
    },
    {
        'title': 'mapel',
        'fields': [
            ('id', 'INT', '11', 'Primary Key, Auto Increment'),
            ('nama_mapel', 'VARCHAR', '100', 'Nama mata pelajaran'),
            ('created_at', 'DATETIME', '-', 'Tanggal dibuat'),
            ('updated_at', 'DATETIME', '-', 'Tanggal update'),
        ]
    },
    {
        'title': 'agama',
        'fields': [
            ('id', 'INT', '11', 'Primary Key, Auto Increment'),
            ('nama_agama', 'VARCHAR', '50', 'Nama agama'),
        ]
    },
    {
        'title': 'guru',
        'fields': [
            ('id', 'INT', '11', 'Primary Key, Auto Increment'),
            ('user_id', 'INT', '11', 'FK ke users'),
            ('nip_nuptk', 'VARCHAR', '50', 'NIP/NUPTK guru'),
            ('nama', 'VARCHAR', '100', 'Nama guru'),
            ('jenis_kelamin', 'ENUM', 'Laki-laki/Perempuan', 'Jenis kelamin'),
            ('agama_id', 'INT', '11', 'FK ke agama'),
            ('tempat_lahir', 'VARCHAR', '100', 'Tempat lahir'),
            ('tanggal_lahir', 'DATE', '-', 'Tanggal lahir'),
            ('mapel_id', 'INT', '11', 'FK ke mapel'),
            ('alamat', 'TEXT', '-', 'Alamat'),
            ('no_hp', 'VARCHAR', '20', 'Nomor HP'),
            ('created_at', 'DATETIME', '-', 'Tanggal dibuat'),
            ('updated_at', 'DATETIME', '-', 'Tanggal update'),
        ]
    },
    {
        'title': 'kelas',
        'fields': [
            ('id', 'INT', '11', 'Primary Key, Auto Increment'),
            ('nama_kelas', 'VARCHAR', '50', 'Nama kelas'),
            ('tingkat', 'INT', '2', 'Tingkat (10/11/12)'),
            ('jurusan_id', 'INT', '11', 'FK ke jurusan'),
            ('wali_kelas_id', 'INT', '11', 'FK ke guru (wali kelas)'),
            ('created_at', 'DATETIME', '-', 'Tanggal dibuat'),
            ('updated_at', 'DATETIME', '-', 'Tanggal update'),
        ]
    },
    {
        'title': 'siswa',
        'fields': [
            ('id', 'INT', '11', 'Primary Key, Auto Increment'),
            ('user_id', 'INT', '11', 'FK ke users'),
            ('nis', 'VARCHAR', '20', 'Nomor Induk Siswa'),
            ('nama', 'VARCHAR', '100', 'Nama siswa'),
            ('tanggal_lahir', 'DATE', '-', 'Tanggal lahir'),
            ('jenis_kelamin', 'ENUM', 'Laki-laki/Perempuan', 'Jenis kelamin'),
            ('agama_id', 'INT', '11', 'FK ke agama'),
            ('kelas_id', 'INT', '11', 'FK ke kelas'),
            ('jurusan_id', 'INT', '11', 'FK ke jurusan'),
            ('alamat', 'TEXT', '-', 'Alamat'),
            ('no_hp', 'VARCHAR', '20', 'Nomor HP'),
            ('created_at', 'DATETIME', '-', 'Tanggal dibuat'),
            ('updated_at', 'DATETIME', '-', 'Tanggal update'),
        ]
    },
    {
        'title': 'jadwal',
        'fields': [
            ('id', 'INT', '11', 'Primary Key, Auto Increment'),
            ('kelas_id', 'INT', '11', 'FK ke kelas'),
            ('mapel_id', 'INT', '11', 'FK ke mapel'),
            ('guru_id', 'INT', '11', 'FK ke guru'),
            ('tahun_akademik_id', 'INT', '11', 'FK ke tahun_akademik'),
            ('hari', 'ENUM', 'Senin-Sabtu', 'Hari pelajaran'),
            ('jam_mulai', 'TIME', '-', 'Jam mulai'),
            ('jam_selesai', 'TIME', '-', 'Jam selesai'),
            ('created_at', 'DATETIME', '-', 'Tanggal dibuat'),
            ('updated_at', 'DATETIME', '-', 'Tanggal update'),
        ]
    },
    {
        'title': 'pertemuan',
        'fields': [
            ('id', 'INT', '11', 'Primary Key, Auto Increment'),
            ('kelas_id', 'INT', '11', 'FK ke kelas'),
            ('mapel_id', 'INT', '11', 'FK ke mapel'),
            ('nama_pertemuan', 'VARCHAR', '100', 'Nama/topik pertemuan'),
            ('tanggal', 'DATE', '-', 'Tanggal pertemuan'),
            ('topik', 'VARCHAR', '255', 'Topik pertemuan'),
            ('video_youtube', 'VARCHAR', '255', 'Link video (opsional)'),
            ('created_at', 'DATETIME', '-', 'Tanggal dibuat'),
            ('updated_at', 'DATETIME', '-', 'Tanggal update'),
        ]
    },
    {
        'title': 'wali_kelas',
        'fields': [
            ('id', 'INT', '11', 'Primary Key, Auto Increment'),
            ('kelas_id', 'INT', '11', 'FK ke kelas'),
            ('guru_id', 'INT', '11', 'FK ke guru'),
            ('tahun_akademik_id', 'INT', '11', 'FK ke tahun_akademik'),
            ('created_at', 'DATETIME', '-', 'Tanggal dibuat'),
            ('updated_at', 'DATETIME', '-', 'Tanggal update'),
        ]
    },
    {
        'title': 'absensi',
        'fields': [
            ('id', 'INT', '11', 'Primary Key, Auto Increment'),
            ('siswa_id', 'INT', '11', 'FK ke siswa'),
            ('tanggal', 'DATE', '-', 'Tanggal absensi'),
            ('status', 'ENUM', 'Hadir/Sakit/Izin/Alpha', 'Status kehadiran'),
            ('created_at', 'DATETIME', '-', 'Tanggal dibuat'),
            ('updated_at', 'DATETIME', '-', 'Tanggal update'),
        ]
    },
    {
        'title': 'nilai',
        'fields': [
            ('id', 'INT', '11', 'Primary Key, Auto Increment'),
            ('siswa_id', 'INT', '11', 'FK ke siswa'),
            ('mapel_id', 'INT', '11', 'FK ke mapel'),
            ('semester', 'ENUM', '1/2', 'Semester'),
            ('uts', 'DECIMAL', '5,2', 'Nilai UTS'),
            ('uas', 'DECIMAL', '5,2', 'Nilai UAS'),
            ('tugas', 'DECIMAL', '5,2', 'Nilai tugas'),
            ('akhir', 'DECIMAL', '5,2', 'Nilai akhir'),
            ('created_at', 'DATETIME', '-', 'Tanggal dibuat'),
            ('updated_at', 'DATETIME', '-', 'Tanggal update'),
        ]
    },
    {
        'title': 'materi',
        'fields': [
            ('id', 'INT', '11', 'Primary Key, Auto Increment'),
            ('guru_id', 'INT', '11', 'FK ke guru'),
            ('mapel_id', 'INT', '11', 'FK ke mapel'),
            ('kelas_id', 'INT', '11', 'FK ke kelas'),
            ('judul', 'VARCHAR', '100', 'Judul materi'),
            ('deskripsi', 'TEXT', '-', 'Deskripsi materi'),
            ('file', 'VARCHAR', '255', 'File materi'),
            ('created_at', 'DATETIME', '-', 'Tanggal dibuat'),
            ('updated_at', 'DATETIME', '-', 'Tanggal update'),
        ]
    },
    {
        'title': 'tugas',
        'fields': [
            ('id', 'INT', '11', 'Primary Key, Auto Increment'),
            ('guru_id', 'INT', '11', 'FK ke guru'),
            ('mapel_id', 'INT', '11', 'FK ke mapel'),
            ('kelas_id', 'INT', '11', 'FK ke kelas'),
            ('deskripsi', 'TEXT', '-', 'Deskripsi tugas'),
            ('deadline', 'DATETIME', '-', 'Deadline tugas'),
            ('file', 'VARCHAR', '255', 'File tugas (opsional)'),
            ('created_at', 'DATETIME', '-', 'Tanggal dibuat'),
            ('updated_at', 'DATETIME', '-', 'Tanggal update'),
        ]
    },
    {
        'title': 'pengumpulan_tugas',
        'fields': [
            ('id', 'INT', '11', 'Primary Key, Auto Increment'),
            ('tugas_id', 'INT', '11', 'FK ke tugas'),
            ('siswa_id', 'INT', '11', 'FK ke siswa'),
            ('file_tugas', 'VARCHAR', '255', 'File tugas yang dikumpulkan'),
            ('status', 'ENUM', 'Belum Dikumpulkan/Sudah Dikumpulkan/Terlambat', 'Status pengumpulan'),
            ('nilai', 'INT', '3', 'Nilai tugas (opsional)'),
            ('catatan', 'TEXT', '-', 'Catatan penilaian'),
            ('created_at', 'DATETIME', '-', 'Tanggal dibuat'),
            ('updated_at', 'DATETIME', '-', 'Tanggal update'),
        ]
    },
    {
        'title': 'pretest',
        'fields': [
            ('id', 'INT', '11', 'Primary Key, Auto Increment'),
            ('pertemuan_id', 'INT', '11', 'FK ke pertemuan'),
            ('judul', 'VARCHAR', '100', 'Judul pretest'),
            ('created_at', 'DATETIME', '-', 'Tanggal dibuat'),
            ('updated_at', 'DATETIME', '-', 'Tanggal update'),
        ]
    },
    {
        'title': 'pretest_soal',
        'fields': [
            ('id', 'INT', '11', 'Primary Key, Auto Increment'),
            ('pretest_id', 'INT', '11', 'FK ke pretest'),
            ('soal', 'TEXT', '-', 'Isi soal'),
            ('pilihan_a', 'VARCHAR', '255', 'Pilihan A'),
            ('pilihan_b', 'VARCHAR', '255', 'Pilihan B'),
            ('pilihan_c', 'VARCHAR', '255', 'Pilihan C'),
            ('pilihan_d', 'VARCHAR', '255', 'Pilihan D'),
            ('jawaban_benar', 'CHAR', '1', 'Kunci jawaban (a/b/c/d)'),
            ('created_at', 'DATETIME', '-', 'Tanggal dibuat'),
            ('updated_at', 'DATETIME', '-', 'Tanggal update'),
        ]
    },
]

for t in tabels:
    doc.add_heading(f"Tabel: {t['title']}", level=2)
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'No'
    hdr_cells[1].text = 'Nama Atribut'
    hdr_cells[2].text = 'Tipe Data'
    hdr_cells[3].text = 'Size'
    hdr_cells[4].text = 'Keterangan'
    for i, f in enumerate(t['fields'], 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = f[0]
        row_cells[2].text = f[1]
        row_cells[3].text = f[2]
        row_cells[4].text = f[3]
    doc.add_paragraph('')

doc.save('struktur_tabel_siakad.docx')
print('File struktur_tabel_siakad.docx berhasil dibuat.') 